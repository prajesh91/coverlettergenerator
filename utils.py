from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import os
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
from pdfminer.high_level import extract_text


def call_llm(prompt, api_key):
    """
    Calls Google Gemini to generate content.
    """
    genai.configure(api_key=api_key)
    try:
        # Using gemini-1.5-flash for better free tier quota limits
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        # Return detailed error for debugging
        return f"Error calling Gemini API: {str(e)}"

def extract_text_from_url(url):
    """
    Extracts text from a job description URL using Jina Reader for better compatibility.
    Includes special handling for LinkedIn URLs to convert them to public viewable links.
    """
    try:
        # Special handling for LinkedIn personalized/collection URLs
        if "linkedin.com" in url and "currentJobId=" in url:
            import re
            match = re.search(r'currentJobId=(\d+)', url)
            if match:
                job_id = match.group(1)
                url = f"https://www.linkedin.com/jobs/view/{job_id}/"

        # Use r.jina.ai to fetch the content as markdown/text
        jina_url = f"https://r.jina.ai/{url}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(jina_url, headers=headers, timeout=15)
        response.raise_for_status()
        
        content = response.text
        
        # Check for common bot-blocking or "not found" indicators in the content
        block_indicators = [
            "Just a moment...",
            "Checking your browser",
            "Access Denied",
            "Page not found",
            "لم يتم العثور على الصفحة", # LinkedIn 404 in Arabic sometimes
            "Direct target URL returned error 403",
            "Direct target URL returned error 404"
        ]
        
        if any(indicator in content for indicator in block_indicators):
            if "error 403" in content or "Just a moment" in content:
                raise Exception("The job board blocked our automated access. Please copy-paste the job description manually.")
            elif "error 404" in content or "not found" in content.lower():
                raise Exception("The job page could not be found. Please check the URL.")
            else:
                raise Exception("We couldn't extract the job details from this link. Please copy-paste it manually.")

        return content
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 403:
            raise Exception("Access blocked by the website. Manual copy-paste is required.")
        elif e.response.status_code == 404:
            raise Exception("Job page not found. Please verify the URL.")
        else:
            raise Exception(f"HTTP error occurred: {e.response.status_code}")
    except Exception as e:
        raise Exception(f"Extraction failed: {str(e)}")

from pdfminer.high_level import extract_text

def extract_text_from_pdf(file):
    """
    Extracts text from a PDF file.
    """
    return extract_text(file)

def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file.
    """
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

import re

def clean_text(text):
    """
    Removes markdown formatting like **bold**, --, etc.
    """
    # Remove bold/italic markers
    text = re.sub(r'\*\*|__', '', text)
    text = re.sub(r'\*|_', '', text)
    # Remove headers
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    # Remove list markers if they are just dashes (optional, but user asked for no --)
    # Keeping bullet points might be good for readability, but user said "no --"
    # Let's just remove double dashes which are often used as separators
    text = re.sub(r'--', '', text)
    return text.strip()

def analyze_ats_score(resume_text, job_description, api_key):
    """
    Analyzes the resume against the job description and provides an ATS score.
    """
    prompt = f"""
    You are an expert ATS (Applicant Tracking System) scanner. 
    Analyze the following resume against the job description.
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Output Format:
    Provide a detailed analysis in plain text. Do NOT use markdown formatting (no bold, no italics, no headers).
    
    Match Score: [Score]/100
    
    Missing Keywords:
    - [Keyword 1]
    - [Keyword 2]
    
    Improvement Suggestions:
    - [Suggestion 1]
    - [Suggestion 2]
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_interview_questions(resume_text, job_description, api_key):
    """
    Generates industry-specific and technical interview questions based on resume and JD.
    """
    prompt = f"""
    You are an expert interviewer specializing in technical and industry-standard evaluations. 
    Based on the candidate's resume and the job description, generate 10 probable industry-specific and technical interview questions.
    
    CRITICAL INSTRUCTIONS:
    1. Output in plain text. No markdown formatting.
    2. Do NOT use double dashes (--).
    3. For EACH question, provide:
       - The Question
       - Why it's being asked
       - An Outline for the Answer (a skeleton of what the candidate should mention)
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_career_insights(resume_text, job_description, api_key):
    """
    Generates career insights including salary negotiation and growth.
    """
    prompt = f"""
    You are a career consultant. Based on the candidate's resume and the job description, provide the following insights:
    1. Salary Negotiation: Estimated range based on industry status and specific tips for this role.
    2. Career Growth: A potential growth chart/pathway for someone in this position.
    3. Outcome of the Job: What the candidate can expect to achieve in terms of skill development and career impact.
    
    CRITICAL INSTRUCTIONS:
    1. Output in plain text. No markdown formatting.
    2. Do NOT use double dashes (--).
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_resume_content(resume_text, job_description, api_key):
    """
    Generates tailored resume content using LLM.
    """
    prompt = f"""
    You are an expert professional resume writer. Rewrite the following resume to tailor it for the job description provided.
    
    CRITICAL INSTRUCTIONS:
    1. Write in a purely human, professional tone. Avoid robotic transitions or overused AI phrases.
    2. Do NOT use any markdown formatting. No bold (**), no italics (*), no headers (#).
    3. Do NOT use double dashes (--).
    4. Target a 90%+ ATS match rate by naturally integrating keywords.
    5. Output ONLY the resume content. No intro/outro.
    
    Job Description:
    {job_description}
    
    Original Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_cover_letter_content(resume_text, job_description, api_key):
    """
    Generates cover letter content using LLM.
    """
    prompt = f"""
    You are an expert career coach. Write a persuasive cover letter based on the candidate's resume and the job description.
    
    CRITICAL INSTRUCTIONS:
    1. Write in a purely human, professional, and engaging tone. 
    2. Avoid generic AI phrases like "I am writing to express my interest". Be more creative and direct.
    3. Do NOT use any markdown formatting. No bold (**), no italics (*).
    4. Do NOT use double dashes (--).
    5. Do not include placeholders like [Your Name] if the information is available in the resume.
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_screening_questions(resume_text, job_description, api_key):
    """
    Generates 5-7 industry-standard screening questions based on JD and Resume.
    """
    prompt = f"""
    You are an expert recruiter. Based on the job description and the candidate's resume, generate 5-7 industry-standard screening questions.
    These should be questions that a recruiter would likely ask during an initial phone screen (e.g., salary expectations, relocation, core skills).
    
    CRITICAL INSTRUCTIONS:
    1. Output in plain text. No markdown formatting.
    2. Do NOT use double dashes (--).
    3. For EACH question, provide:
       - The Question
       - A brief tip on why they are asking
       - An Outline for the Answer (suggested content based on their resume)
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_final_interview_questions(resume_text, job_description, api_key):
    """
    Generates final-round behavioral and culture-fit questions.
    """
    prompt = f"""
    You are a Hiring Manager preparing for a final-round interview. 
    Based on the candidate's resume and the job description, generate 5 high-impact final interview questions.
    Focus on long-term fit, behavioral scenarios, and executive presence.
    
    CRITICAL INSTRUCTIONS:
    1. Output in plain text. No markdown formatting.
    2. Do NOT use double dashes (--).
    3. For EACH question, provide:
       - The Question
       - The underlying trait being tested
       - An Outline for the Answer (recommended structure for a winning response)
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def provide_interview_feedback(question, answer, job_description, api_key):
    """
    Provides feedback on a user's answer to a specific interview question.
    """
    prompt = f"""
    You are an expert interview coach. Evaluate the following answer to an interview question.
    
    Job Description:
    {job_description}
    
    Question:
    {question}
    
    Candidate's Answer:
    {answer}
    
    Output Format:
    - Feedback: [Detailed feedback on the strengths and weaknesses of the answer]
    - Suggestion: [How to improve the answer using the STAR method if applicable]
    - Improved Answer: [A sample of how a strong candidate would answer this]
    
    CRITICAL INSTRUCTIONS:
    1. Output in plain text. No markdown formatting.
    2. Do NOT use double dashes (--).
    """
    response = call_llm(prompt, api_key)
    return clean_text(response)

def generate_docx_from_text(text_content):
    """
    Generates a DOCX file from raw text.
    """
    doc = Document()
    # Split by newlines and add as paragraphs
    for line in text_content.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf_from_text(text_content):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40
    
    # Simple text wrapping could be added here
    for line in text_content.split('\n'):
        if y < 40:
            p.showPage()
            y = height - 40
        p.drawString(40, y, line[:90]) # Simple truncation to avoid overflow
        y -= 15
        
    p.save()
    buffer.seek(0)
    return buffer
