from docx import Document

def create_sample_resume():
    path = "/Users/prajeshpoudyal/Desktop/cover letter/sample_resume.docx"
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer | Python Expert')
    doc.add_paragraph('Email: john.doe@example.com')
    
    doc.add_heading('Experience', level=1)
    p = doc.add_paragraph('Senior Developer at Tech Corp (2020 - Present)')
    p.add_run('\n- Led development of AI-driven tools.')
    p.add_run('\n- Optimized database queries by 40%.')
    
    doc.save(path)
    print(f"Sample resume created at {path}")

if __name__ == "__main__":
    create_sample_resume()
